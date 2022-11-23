import docx

# doc = docx.Document()
# doc.save('customInvites.docx')

txt = open('../automate_online-materials2e/guests.txt')
names = txt.readlines()
names = [x.split('\n')[0] for x in names]
txt.close()

doc = docx.Document('customInvites.docx')

i = 1
for name in names:
    doc.add_paragraph('It would be a pleasure to have the company of', style='Invites')
    doc.add_paragraph(name, style='Name')
    doc.add_paragraph('at 11010 Memory Lane on the Evening of', style='Invites')
    doc.add_paragraph('April 1st', style='DateLine')
    last_line = doc.add_paragraph('at 7 o\'clock', style='Invites')
    # doc.paragraphs[i].runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)
    last_line.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

doc.save('customInvites_final.docx')
print('File has been saved as customInvites_final.docx')